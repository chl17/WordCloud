# .docx
import docx
# .doc
from win32com import client
# .pdf
import subprocess
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator
# .xls and .xlsx
import xlrd
# timeout
import func_timeout
from func_timeout import func_set_timeout
# pdf2txt.py
import os
from src.osdir import all_path


def readattachment(file_paths):
    tempfile = open(os.getcwd() + '\\temp.txt', mode="w+", encoding='utf-8')
    for file_path in file_paths:
        print("Reading: " + file_path)
        # .pdf
        if file_path.endswith('.pdf'):
            """
            try:  # 防止读取 PDF 超时阻塞 使用 pdf2txt.py
                content = pdf2text(file_path)
                if content.strip():  # 排除无文字读取的 PDF
                    content_all.append(content)
            except func_timeout.exceptions.FunctionTimedOut:
                content_all.append(file_path + ' 读取超时!')
            """
            try:  # 防止读取 PDF 超时阻塞 使用 PDFMiner
                content = readPDF(file_path)
                tempfile.writelines(content)
            except func_timeout.exceptions.FunctionTimedOut:
                print(file_path + ' 读取超时!')
        # .doc
        if file_path.endswith('.doc'):
            doc2docx(file_path)
            doc = docx.Document(file_path + 'x')
            content = []
            for para in doc.paragraphs:
                content.append(para.text)
            tempfile.writelines(content)
        # .docx
        if file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            content = []
            for para in doc.paragraphs:
                content.append(para.text)
            tempfile.writelines(content)
        # .xls and .xlsx
        if file_path.endswith('.xls') or file_path.endswith('.xlsx'):
            data = xlrd.open_workbook(file_path)
            table_names = data.sheet_names()
            content_list = []
            content_cleaned = []
            for table_name in table_names:
                table = data.sheet_by_name(table_name)
                col_count = table.ncols
                for col in range(col_count):
                    content_list = content_list + table.col_values(col)
            for element in content_list:
                if element != '' and (not ((type(element) == int) or (type(element) == float))):
                    content_cleaned.append(str(element))
            tempfile.writelines(content_cleaned)


def doc2docx(doc_name):
    docx_name = doc_name + "x"
    if not os.path.exists(docx_name):
        print("Converting " + doc_name + " to docx")
        # 首先将doc转换成docx
        word = client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_name)
        #使用参数16表示将doc转换成docx
        doc.SaveAs(docx_name, 16)
        doc.Close()
        word.Quit()


@func_set_timeout(5)  # 设定超时限制5s
def readPDF(file_path):
    fp = open(file_path, 'rb')
    # 来创建一个pdf文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档对象存储文档结构
    document = PDFDocument(parser)
    # 检查文件是否允许文本提取
    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建一个PDF资源管理器对象来存储共赏资源
        rsrcmgr = PDFResourceManager()
        # 设定参数进行分析
        laparams = LAParams()
        # 创建一个PDF设备对象
        # device=PDFDevice(rsrcmgr)
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        content = []
        # 处理每一页
        for page in PDFPage.create_pages(document):
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    text_stripped = x.get_text().strip()
                    if text_stripped != '':
                        content.append(x.get_text().strip())

        return '\n'.join(content)


@func_set_timeout(5)  # 设定超时限制5s  http://www.cnblogs.com/hester/p/7641258.html
def pdf2text(file_path):
    cmd = ['python', os.getcwd() + "\\pdf2txt.py", file_path]
    output_text = subprocess.check_output(cmd).decode("utf-8")
    return output_text


if __name__ == '__main__':
    filePath = r"C:\Users\haoli\Documents\国家能源互联网政策库\地方政策\河北省\1、能源供给革命政策"
    filePaths = all_path(filePath)
    readattachment(filePaths)
    pass
